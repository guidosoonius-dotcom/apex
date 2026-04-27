"""
/api/generate-docx — Convert markdown deliverable to branded Word document.
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from http.server import BaseHTTPRequestHandler
from _shared import CORS_HEADERS, send_cors_ok, read_json, parse_md, strip_inline, is_meta, slug
import json
import io

NAVY  = (15,  27,  45)
TEAL  = (14, 124, 123)
GOLD  = (201, 168,  76)
WHITE = (240, 245, 250)
FOG   = (139, 163, 190)


def build_docx(content, title, client, sector):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    def rgb(*t): return RGBColor(*t)

    def run(para, text, bold=False, color=None, size=None, italic=False):
        r = para.add_run(text)
        r.bold   = bold
        r.italic = italic
        if color: r.font.color.rgb = rgb(*color)
        if size:  r.font.size = Pt(size)
        return r

    def inline(para, text, size=11, default_color=None):
        import re
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for i, part in enumerate(parts):
            if part:
                run(para, part, bold=(i % 2 == 1), size=size, color=default_color)

    cover = doc.add_paragraph()
    cover.alignment = 1
    run(cover, title, bold=True, color=TEAL, size=28)

    sub = doc.add_paragraph()
    sub.alignment = 1
    run(sub, '  ·  '.join(filter(None, [client, sector])), color=FOG, size=12)
    doc.add_paragraph()

    for btype, btext in parse_md(content):
        if btype == 'title':
            pass
        elif btype == 'h1':
            h = doc.add_heading(level=1); h.clear()
            run(h, strip_inline(btext), bold=True, color=TEAL, size=16)
        elif btype == 'h2':
            h = doc.add_heading(level=2); h.clear()
            run(h, strip_inline(btext), bold=True, color=GOLD, size=13)
        elif btype == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            inline(p, strip_inline(btext), size=11)
        elif btype == 'hr':
            doc.add_paragraph()
        elif btype == 'para':
            if is_meta(btext): continue
            p = doc.add_paragraph()
            inline(p, strip_inline(btext), size=11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class handler(BaseHTTPRequestHandler):

    def log_message(self, *a): pass

    def do_OPTIONS(self):
        send_cors_ok(self)

    def do_POST(self):
        try:
            p       = read_json(self)
            content = p.get('content', '')
            title   = p.get('title',   'APEX Deliverable')
            client  = p.get('client',  '')
            sector  = p.get('sector',  '')
            data     = build_docx(content, title, client, sector)
            mime     = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            filename = slug(title) + '.docx'
            self.send_response(200)
            for k, v in CORS_HEADERS.items():
                self.send_header(k, v)
            self.send_header('Content-Type',        mime)
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
            self.send_header('Content-Length',      str(len(data)))
            self.end_headers()
            self.wfile.write(data)
        except Exception as e:
            import traceback; traceback.print_exc()
            msg = json.dumps({'error': str(e)}).encode()
            self.send_response(500)
            for k, v in CORS_HEADERS.items():
                self.send_header(k, v)
            self.send_header('Content-Type',   'application/json')
            self.send_header('Content-Length', str(len(msg)))
            self.end_headers()
            self.wfile.write(msg)
